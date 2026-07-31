"""
Tests for the Purchase Order Word (.docx) generator's Bank Details box.

Constraint #9: generation happens in memory — these tests read the generated
bytes straight back with python-docx, never touching disk.
"""
import pytest
from io import BytesIO

from docx import Document

from apps.master_data.tests.factories import BankFactory
from pdf.purchase_order_word import generate_po_docx

from .factories import PurchaseOrderFactory, PurchaseOrderLineItemFactory


def _docx_text(po) -> str:
    document = Document(BytesIO(generate_po_docx(po).getvalue()))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.django_db
class TestPurchaseOrderWordBankDetails:
    def test_ad_code_and_lut_number_appear_in_bank_box(self):
        bank = BankFactory(ad_code="14012345678901", lut_number="AD070124000123A", lut_valid_until="2027-03-31")
        po = PurchaseOrderFactory(bank=bank)
        PurchaseOrderLineItemFactory(purchase_order=po)
        text = _docx_text(po)
        assert "14012345678901" in text
        assert "AD070124000123A" in text

    def test_blank_ad_code_and_lut_omitted(self):
        bank = BankFactory(ad_code="", lut_number="")
        po = PurchaseOrderFactory(bank=bank)
        PurchaseOrderLineItemFactory(purchase_order=po)
        text = _docx_text(po)
        assert "AD CODE" not in text.upper()
        assert "LUT" not in text.upper()
