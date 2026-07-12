"""
Proforma Invoice Word (.docx) Generator

Transliterates pdf/proforma_invoice.py's actual ReportLab Table()/TableStyle()
calls into python-docx tables, cell-by-cell and span-by-span, using the
docx_base primitives (build_grid_table / add_title_block / add_html_runs /
build_items_table / build_lined_box). No repeating per-page header — the
company name and doc title are ordinary body content printed once, exactly
like pdf/proforma_invoice.py's `story.append(Paragraph(...))` calls. Only the
footer + DRAFT watermark repeat on every page (via section header/footer),
matching the PDF's NumberedCanvas.
"""
import html
import io
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Mm

from pdf.docx_base import (
    CONTENT_W,
    add_page_footer,
    add_run,
    add_title_block,
    add_watermark,
    build_grid_table as _build_grid_table,
    build_items_table,
    build_lined_box as _build_lined_box,
    setup_page,
)
from pdf.proforma_invoice import (
    _org_address_str,
    _org_country_name,
    _org_email,
    amount_to_words,
    bool_yn,
    fmt_date,
    fmt_money,
    fmt_qty,
    fmt_rate,
    safe,
)


def _unescape_html_entities(spec):
    """
    The ported ReportLab strings use HTML entities (&amp;, &nbsp;, etc.) because
    ReportLab's Paragraph is a mini HTML/XML renderer. docx_base's add_html_runs
    only understands <b>/</b>/<br/> tags — it does not decode entities — so they
    must be unescaped here before being handed to build_grid_table, otherwise
    "&amp;" renders literally instead of as "&".
    """
    if spec is None:
        return None
    if isinstance(spec, str):
        return html.unescape(spec)
    if isinstance(spec, dict) and "html" in spec:
        spec = dict(spec)
        spec["html"] = html.unescape(spec["html"])
    return spec


def build_grid_table(document, rows, col_widths, spans=None):
    """Thin wrapper around docx_base.build_grid_table that unescapes HTML
    entities in every cell before rendering (see _unescape_html_entities)."""
    clean_rows = [[_unescape_html_entities(cell) for cell in row] for row in rows]
    return _build_grid_table(document, clean_rows, col_widths, spans=spans)


def build_lined_box(document, lines_html, bg_hex=None, size=None):
    """Thin wrapper around docx_base.build_lined_box that unescapes HTML
    entities in every line before rendering (see _unescape_html_entities)."""
    clean_lines = [html.unescape(line) for line in lines_html]
    kwargs = {}
    if bg_hex is not None:
        kwargs["bg_hex"] = bg_hex
    if size is not None:
        kwargs["size"] = size
    return _build_lined_box(document, clean_lines, **kwargs)


def generate_proforma_invoice_docx_bytes(invoice, client_invoice=False) -> bytes:
    """
    Generate Proforma Invoice Word document, transliterating the ReportLab
    Table()/TableStyle() layout from generate_proforma_invoice_pdf_bytes()
    section-by-section.
    """
    from apps.workflow.constants import APPROVED

    document = Document()
    setup_page(document)
    section = document.sections[0]

    is_draft = getattr(invoice, "status", None) != APPROVED

    # Footer repeats every page (matches NumberedCanvas._draw_footer's
    # "Page N of Total" line).
    add_page_footer(section, with_total=True)
    if is_draft:
        add_watermark(document, text="DRAFT")

    # ========================================================================
    # SECTION 1: DOCUMENT HEADER — one-time title block, NOT a repeating
    # per-page header (mirrors story.append(Paragraph(exp_name, ...)) then
    # story.append(Paragraph(doc_title, ...)) plus the LINEABOVE separator).
    # ========================================================================
    exp = invoice.exporter
    exp_name = safe(getattr(exp, "name", ""))
    exp_address = _org_address_str(exp)
    exp_email = _org_email(exp)
    _exp_addr_obj = exp.addresses.first() if exp else None
    exp_iec = safe(getattr(_exp_addr_obj, "iec_code", ""))
    exp_country = _org_country_name(exp)

    exp_detail_parts = []
    if exp_address:
        exp_detail_parts.append(exp_address)
    if exp_country:
        exp_detail_parts.append(exp_country)
    if exp_email:
        exp_detail_parts.append(exp_email)
    exp_detail_html = "<br/>".join([exp_name] + exp_detail_parts)

    doc_title = "PROFORMA INVOICE CUM SALES CONTRACT"
    add_title_block(document, exp_name, [doc_title])

    # ========================================================================
    # SECTION 2: MAIN INFO TABLE (TOP HALF) — main_info_data / main_info_table_top
    # 4 cols x 4 rows. ReportLab's SPAN command uses (col,row) coordinates:
    # ("SPAN", (start_col,start_row), (end_col,end_row)). Verified empirically
    # by rendering an isolated repro table and inspecting the resulting PDF:
    #   SPAN (0,0)-(1,1)  -> cols 0-1, rows 0-1  => Exporter block (2x2 merge)
    #   SPAN (0,2)-(1,3)  -> cols 0-1, rows 2-3  => Consignee block (2x2 merge)
    #   SPAN (2,2)-(2,3)  -> col 2,   rows 2-3   => Buyer block (vertical merge)
    #   SPAN (3,2)-(3,3)  -> col 3,   rows 2-3   => Other reference(s) block (vertical merge)
    # Row0/Row1 in cols 2-3 are NOT merged — PI No&Date / Buyer Order No&Date
    # sit directly above Country of Origin / Country of Final Destination as
    # separate cells. build_grid_table's spans are (r1,c1,r2,c2) tuples.
    # ========================================================================
    cons = invoice.consignee
    cons_name = safe(getattr(cons, "name", ""))
    cons_address = _org_address_str(cons)
    cons_email = _org_email(cons)
    cons_country = _org_country_name(cons)

    address_parts = []
    if cons_address:
        address_parts.append(cons_address)
    if cons_country:
        address_parts.append(cons_country)
    address_line = ", ".join(address_parts)

    cons_lines = [x for x in [cons_name, address_line, cons_email] if x]
    consignee_details_html = "<br/>".join(cons_lines)

    buyer_obj = getattr(invoice, "buyer", None)
    buyer_name = safe(getattr(buyer_obj, "name", "")) if buyer_obj else cons_name

    if buyer_obj:
        buyer_address = _org_address_str(buyer_obj)
        buyer_email = _org_email(buyer_obj)
        buyer_country = _org_country_name(buyer_obj)
        buyer_addr_parts = []
        if buyer_address:
            buyer_addr_parts.append(buyer_address)
        if buyer_country:
            buyer_addr_parts.append(buyer_country)
        buyer_addr_line = ", ".join(buyer_addr_parts)
        buyer_lines = [x for x in [buyer_name, buyer_addr_line, buyer_email] if x]
        buyer_details_html = "<br/>".join(buyer_lines)
    else:
        buyer_details_html = cons_name

    origin_country = _org_country_name(exp)
    final_dest_obj = getattr(invoice, "country_of_final_destination", None)
    final_country = safe(getattr(final_dest_obj, "name", "")) if final_dest_obj else ""

    pi_number = safe(getattr(invoice, "pi_number", ""))
    pi_date = fmt_date(getattr(invoice, "pi_date", None))
    buyer_order_no = safe(getattr(invoice, "buyer_order_no", ""))
    buyer_order_date = fmt_date(getattr(invoice, "buyer_order_date", None))
    other_references = safe(getattr(invoice, "other_references", ""))

    # ReportLab's SPAN command is (col_start,row_start),(col_end,row_end).
    # ("SPAN", (0, 0), (1, 1))  -> cols 0-1, rows 0-1  => Exporter block
    # ("SPAN", (0, 2), (1, 3))  -> cols 0-1, rows 2-3  => Consignee block
    # ("SPAN", (2, 2), (2, 3))  -> col 2,   rows 2-3   => Buyer block
    # ("SPAN", (3, 2), (3, 3))  -> col 3,   rows 2-3   => Other reference(s) block
    # build_grid_table's spans are (r1,c1,r2,c2) — converted accordingly below.
    col_w4 = Emu(int(CONTENT_W / 4))
    main_info_rows = [
        [
            {"html": f"<b>Exporter:</b><br/>{exp_detail_html}"},
            None,
            {"html": f"<b>Proforma Invoice No &amp; Date:</b><br/>{pi_number} &amp; {pi_date}"},
            {"html": f"<b>Buyer Order No and Date:</b><br/>{buyer_order_no} &amp; {buyer_order_date}"},
        ],
        [
            None, None,
            {"html": f"<b>Country of Origin of Goods:</b><br/>{origin_country}"},
            {"html": f"<b>Country of Final Destination:</b><br/>{final_country}"},
        ],
        [
            {"html": f"<b>Consignee:</b><br/>{consignee_details_html}"},
            None,
            {"html": f"<b>Buyer if other than consignee:</b><br/>{buyer_details_html}"},
            {"html": f"<b>Other reference(s):</b><br/>{other_references}"},
        ],
        [None, None, None, None],
    ]
    build_grid_table(
        document, main_info_rows, [col_w4] * 4,
        # (r1,c1,r2,c2), matching the four verified SPAN commands above.
        spans=[(0, 0, 1, 1), (2, 0, 3, 1), (2, 2, 3, 2), (2, 3, 3, 3)],
    )
    document.add_paragraph()

    # ========================================================================
    # SECTION 3: MAIN INFO TABLE (BOTTOM HALF) — info_rows_b (3x2) + info_row_a (4x1)
    # ========================================================================
    incoterm_obj = getattr(invoice, "incoterms", None)
    incoterm_disp = safe(getattr(incoterm_obj, "code", "")) if incoterm_obj else ""

    payment_term_obj = getattr(invoice, "payment_terms", None)
    payment_term_name = safe(getattr(payment_term_obj, "name", "")) if payment_term_obj else ""

    port_of_loading_obj = getattr(invoice, "port_of_loading", None)
    port_loading = safe(getattr(port_of_loading_obj, "name", "")) if port_of_loading_obj else ""

    port_of_discharge_obj = getattr(invoice, "port_of_discharge", None)
    port_discharge = safe(getattr(port_of_discharge_obj, "name", "")) if port_of_discharge_obj else ""

    final_dest_location_obj = getattr(invoice, "final_destination", None)
    final_dest = safe(getattr(final_dest_location_obj, "name", "")) if final_dest_location_obj else ""

    pre_carriage_obj = getattr(invoice, "pre_carriage_by", None)
    pre_carriage = safe(getattr(pre_carriage_obj, "name", "")) if pre_carriage_obj else ""

    por_pre_obj = getattr(invoice, "place_of_receipt_by_pre_carrier", None)
    por_pre = safe(getattr(por_pre_obj, "name", "")) if por_pre_obj else ""

    col_w3 = Emu(int(CONTENT_W / 3))
    build_grid_table(document, [
        [
            {"html": f"<b>Pre-Carriaged By:</b><br/>{pre_carriage}"},
            {"html": f"<b>Place of Receipt by Pre-Carrier:</b><br/>{por_pre}"},
            {"html": f"<b>Vessel/Flight No:</b><br/>{safe(invoice.vessel_flight_no)}"},
        ],
        [
            {"html": f"<b>No &amp; Kind of Packages</b><br/>{safe(invoice.kind_of_packages)}"},
            {"html": f"<b>Marks &amp; Nos/Container No</b><br/>{safe(invoice.marks_and_nos)}"},
            {"html": f"<b>Payment Terms:</b><br/>{payment_term_name}"},
        ],
    ], [col_w3] * 3)

    build_grid_table(document, [[
        {"html": f"<b>Port of Loading:</b><br/>{port_loading}"},
        {"html": f"<b>Port of Discharge:</b><br/>{port_discharge}"},
        {"html": f"<b>Final Destination:</b><br/>{final_dest}"},
        {"html": f"<b>Incoterms:</b><br/>{incoterm_disp}"},
    ]], [col_w4] * 4)
    document.add_paragraph()

    # ========================================================================
    # SECTION 4: LINE ITEMS TABLE — li_table (navy header, right-aligned Qty/Rate/Amount)
    # ========================================================================
    currency_code = invoice.currency.code

    line_items_total = Decimal("0.00")

    # For client invoice, pre-compute CIF-adjusted rate and amount per line item.
    # Same formula used in the UI's renderCIFRateCalculation component and in
    # generate_proforma_invoice_pdf_bytes().
    cif_overrides = {}  # it.pk → (cif_rate: float, cif_amount: Decimal)
    if client_invoice:
        all_items = list(invoice.line_items.all().order_by("id"))
        total_qty_all = sum(float(getattr(it, "quantity", 0) or 0) for it in all_items)
        total_fob_all = sum(float(getattr(it, "amount", 0) or 0) for it in all_items)
        total_freight = float(getattr(invoice, "freight", None) or 0)
        total_insurance = float(getattr(invoice, "insurance_amount", None) or 0)
        freight_per_unit = total_freight / total_qty_all if total_qty_all > 0 else 0
        for it in all_items:
            base_rate = float(getattr(it, "rate", None) or 0)
            qty = float(getattr(it, "quantity", None) or 0)
            ins_per_unit = (total_insurance * base_rate) / total_fob_all if total_fob_all > 0 else 0
            cif_rate = base_rate + freight_per_unit + ins_per_unit
            cif_overrides[it.pk] = (cif_rate, Decimal(str(round(cif_rate * qty, 2))))

    li_rows = []
    for idx, it in enumerate(invoice.line_items.all().order_by("id"), start=1):
        uom_obj = getattr(it, "uom", None)
        uom_display = safe(getattr(uom_obj, "abbreviation", "")) if uom_obj else ""
        if client_invoice and it.pk in cif_overrides:
            display_rate, amount = cif_overrides[it.pk]
        else:
            display_rate = getattr(it, "rate", None)
            amount = getattr(it, "amount", None)
        if amount is not None:
            try:
                line_items_total += Decimal(str(amount))
            except Exception:
                pass
        li_rows.append([
            str(idx),
            safe(it.hsn_code),
            safe(it.item_code),
            safe(it.description),
            f"{fmt_qty(it.quantity)} {uom_display}".strip(),
            fmt_rate(display_rate),
            fmt_money(amount),
        ])

    # colWidths=[10, 24, 24, 40, 27, 31, 24] mm, total 180mm — same proportions as the PDF.
    li_col_widths = [Mm(10), Mm(24), Mm(24), Mm(40), Mm(27), Mm(31), Mm(24)]
    build_items_table(
        document,
        headers=["Sr.", "HSN Code", "Item Code", "Description of Goods", "Qty",
                 f"Rate ({currency_code})", f"Amount ({currency_code})"],
        rows=li_rows,
        col_widths=li_col_widths,
        right_cols=[4, 5, 6],
    )
    document.add_paragraph()

    # ========================================================================
    # SECTION 5: TOTALS SECTION — totals_table (2 cols: 140mm/40mm), navy last row
    # ========================================================================
    charges_list = list(invoice.charges.all().order_by("id"))
    charges_total = Decimal("0.00")
    for charge in charges_list:
        charges_total += Decimal(str(charge.amount or 0))

    grand_total = line_items_total + charges_total

    _SELLER_FIELDS = {
        "EXW": [],
        "FCA": [],
        "FOB": [],
        "CFR": ["freight"],
        "CPT": ["freight"],
        "CIF": ["freight", "insurance_amount"],
        "CIP": ["freight", "insurance_amount"],
        "DAP": ["freight", "insurance_amount"],
        "DPU": ["freight", "insurance_amount", "destination_charges"],
        "DDP": ["freight", "insurance_amount", "import_duty", "destination_charges"],
    }
    _FIELD_LABELS = {
        "freight": "Freight",
        "insurance_amount": "Insurance Amount",
        "import_duty": "Import Duty",
        "destination_charges": "Destination Charges",
    }
    seller_fields = _SELLER_FIELDS.get(incoterm_disp, [])
    # For client invoice, CIF rates already include freight & insurance — skip the cost breakdown.
    show_cost_breakdown = bool(incoterm_disp) and incoterm_disp != "EXW" and not client_invoice

    invoice_total_pdf = grand_total
    for field in seller_fields:
        val = getattr(invoice, field, None)
        if val is not None:
            try:
                invoice_total_pdf += Decimal(str(val))
            except Exception:
                pass

    # Client invoice total is just the sum of CIF line amounts (no separate freight/ins rows).
    final_total = grand_total if client_invoice else (invoice_total_pdf if incoterm_disp else grand_total)

    totals_rows = []

    if charges_list and not show_cost_breakdown:
        totals_rows.append([
            {"html": "Item Total"},
            {"html": f"{currency_code} {fmt_money(line_items_total)}", "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ])
        for charge in charges_list:
            totals_rows.append([
                {"html": safe(charge.description)},
                {"html": f"{currency_code} {fmt_money(charge.amount)}", "align": WD_ALIGN_PARAGRAPH.RIGHT},
            ])

    if not incoterm_disp:
        totals_rows.append([
            {"html": "<b>Grand Total Amount</b>", "bg": "navy"},
            {"html": f"<b>{currency_code} {fmt_money(grand_total)}</b>", "bg": "navy", "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ])

    if show_cost_breakdown:
        totals_rows.append([
            {"html": f"<b>Cost Breakdown ({incoterm_disp})</b>"},
            {"html": ""},
        ])
        totals_rows.append([
            {"html": "FOB Value"},
            {"html": f"{currency_code} {fmt_money(grand_total)}", "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ])
        for field in seller_fields:
            val = getattr(invoice, field, None)
            # Skip freight/insurance rows if the user has not filled in a value
            if val is None:
                continue
            totals_rows.append([
                {"html": _FIELD_LABELS.get(field, field)},
                {"html": f"{currency_code} {fmt_money(val)}", "align": WD_ALIGN_PARAGRAPH.RIGHT},
            ])

    if incoterm_disp and client_invoice:
        totals_rows.append([
            {"html": "<b>Total CIF Amount (Payable)</b>", "bg": "navy"},
            {"html": f"<b>{currency_code} {fmt_money(grand_total)}</b>", "bg": "navy", "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ])
    elif incoterm_disp:
        totals_rows.append([
            {"html": "<b>Invoice Total (Amount Payable)</b>", "bg": "navy"},
            {"html": f"<b>{currency_code} {fmt_money(invoice_total_pdf)}</b>", "bg": "navy", "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ])

    if totals_rows:
        build_grid_table(document, totals_rows, [Mm(140), Mm(40)])
        document.add_paragraph()

    # ========================================================================
    # SECTION 6: AMOUNT IN WORDS — words_table (single cell, full width)
    # ========================================================================
    build_grid_table(document, [[
        {"html": f"<b>Amount in Words:</b> {amount_to_words(final_total, currency=currency_code)}"},
    ]], [Mm(180)])
    document.add_paragraph()

    # ========================================================================
    # SECTION 7: VALIDITY & SHIPMENT TABLE — validity_table (2x2, 90mm cols)
    # ========================================================================
    build_grid_table(document, [
        [
            {"html": f"<b>Validity for Acceptance:</b> {fmt_date(invoice.validity_for_acceptance)}"},
            {"html": f"<b>Validity for Shipment:</b> {fmt_date(invoice.validity_for_shipment)}"},
        ],
        [
            {"html": f"<b>Partial Shipment:</b> {bool_yn(invoice.partial_shipment)}"},
            {"html": f"<b>Transshipment:</b> {bool_yn(invoice.transshipment)}"},
        ],
    ], [Mm(90), Mm(90)])
    document.add_paragraph()

    # ========================================================================
    # SECTION 8: DECLARATION — decl_table (single cell, full width)
    # ========================================================================
    decl_html = (
        "<b>Declaration:</b> We declare that this invoice shows the actual price of the "
        "goods described and that all particulars are true and correct."
    )
    if getattr(invoice, "bank_charges_to_buyer", False):
        decl_html += (
            "<br/><b>BANK CHARGES:</b> WITHIN INDIA ON ACCOUNT OF BENEFICIARY &amp; "
            "OUTSIDE OF INDIA ON ACCOUNT OF BUYER"
        )
    build_grid_table(document, [[{"html": decl_html}]], [Mm(180)])
    document.add_paragraph()

    # ========================================================================
    # SECTION 9: BANK DETAILS — build_lined_box (plain white, one row per line)
    # Transliterated verbatim from bank_lines.append(...) in
    # generate_proforma_invoice_pdf_bytes() lines 766-796: same field-label
    # wording/capitalization, same intermediary-institution line, same closing
    # MT103 sentence baked into the box (not a separate paragraph).
    # ========================================================================
    bank = getattr(invoice, "bank", None)
    if bank:
        bank_lines = []
        bank_lines.append(f"<b>BENEFICIARY NAME:</b> {safe(bank.beneficiary_name)}")
        bank_lines.append(f"<b>BANK NAME:</b> {safe(bank.bank_name)}")
        bank_lines.append(f"<b>BRANCH NAME:</b> {safe(bank.branch_name)}")
        bank_lines.append(f"<b>BRANCH ADDRESS:</b> {safe(bank.branch_address)}")
        bank_lines.append(f"<b>A/C NO.:</b> {safe(bank.account_number)}")
        if bank.routing_number:
            bank_lines.append(f"<b>IFSC CODE:</b> {safe(bank.routing_number)}")
        if bank.swift_code:
            bank_lines.append(f"<b>SWIFT CODE:</b> {safe(bank.swift_code)}")
        if bank.iban:
            bank_lines.append(f"<b>IBAN:</b> {safe(bank.iban)}")

        if safe(bank.intermediary_bank_name):
            intermediary_currency_code = (
                safe(getattr(bank.intermediary_currency, "code", ""))
                if getattr(bank, "intermediary_currency", None) else ""
            )
            bank_lines.append(
                f"<b>Intermediary Institution Routing for Currency</b> {intermediary_currency_code} "
                f"<b>A/C No.:</b> {safe(bank.intermediary_account_number)} "
                f"The Bank of {safe(bank.intermediary_bank_name)} "
                f"<b>SWIFT Code:</b> {safe(bank.intermediary_swift_code)}"
            )

        bank_lines.append(
            "Request your bank to send MT 103 Message to our bank and send us copy of this "
            "message to trace &amp; claim the payment from our bank."
        )

        build_lined_box(document, bank_lines)
        document.add_paragraph()

    # ========================================================================
    # SECTION 10: TERMS & CONDITIONS — page break + navy tc_header bar
    # ========================================================================
    tc_content = getattr(invoice, "tc_content", "") or ""
    if tc_content.strip():
        document.add_page_break()
        build_grid_table(document, [[{"html": "<b>Additional Terms &amp; Conditions</b>", "bg": "navy"}]], [Mm(180)])
        document.add_paragraph()

        # tc_content is stored as HTML from a rich text editor; render as plain
        # text paragraphs (the Word equivalent of html_to_rl_flowables's
        # plain-text fallback rendering used by the PDF). Entities (&amp; etc.)
        # must be unescaped same as everywhere else in this file.
        import re as _re
        text_only = _re.sub(r"<[^>]+>", "\n", tc_content)
        for line in [l.strip() for l in text_only.split("\n")]:
            if line:
                p_tc = document.add_paragraph()
                add_run(p_tc, html.unescape(line))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_pi_docx(pi, client_invoice=False) -> io.BytesIO:
    """Wrapper used by the PI view — returns an in-memory BytesIO buffer."""
    docx_bytes = generate_proforma_invoice_docx_bytes(pi, client_invoice=client_invoice)
    buffer = io.BytesIO(docx_bytes)
    buffer.seek(0)
    return buffer
