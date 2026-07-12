"""
Certificate of Analysis Word (.docx) Generator
Transliterates pdf/certificate_of_analysis.py's ReportLab Table()/TableStyle()
calls cell-by-cell so the Word output is visually consistent with the real PDF
(navy header bars, Arial/Helvetica-equivalent body text, one-time title block,
footer-only repetition, DRAFT watermark). Built in memory and streamed — never
written to disk (Rule #9).
"""
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Mm

from pdf.certificate_of_analysis import _safe, _esc, _fmt_date, _fmt_datetime
from pdf.docx_base import (
    setup_page,
    add_page_footer,
    add_watermark,
    add_run,
    set_cell_html,
    set_table_borders,
    set_column_widths,
    build_grid_table,
    _set_cell_background,
    COLOR_NAVY_HEX,
    COLOR_WHITE,
    COLOR_TEXT_DARK,
    FONT_HEADING,
    FONT_BODY,
    SIZE_BODY,
    SIZE_SMALL,
    SIZE_TABLE_HEADER,
    CONTENT_W,
)

# Matches LIGHT_GREY = colors.HexColor('#F0F0F0') in pdf/certificate_of_analysis.py
LIGHT_GREY_HEX = "F0F0F0"


def generate_coa_docx(coa) -> BytesIO:
    """
    Generate a COA Word document for the given CertificateOfAnalysis instance.
    Returns a BytesIO buffer ready to stream — never written to disk.
    """
    from apps.workflow.constants import APPROVED

    document = Document()
    setup_page(document)
    section = document.sections[0]

    # -------------------------------------------------------------------------
    # Gather org / address data — identical logic to generate_coa_pdf()
    # -------------------------------------------------------------------------
    org = coa.footer_organisation
    org_name = _safe(org.name) if org else ""

    org_addresses = []
    org_iec_code = ""
    org_tax_info = ""
    if org:
        for addr in org.addresses.select_related("country").all():
            parts = [addr.line1]
            if addr.line2:
                parts.append(addr.line2)
            parts.append(addr.city)
            if addr.state:
                parts.append(addr.state)
            if addr.pin:
                parts.append(addr.pin)
            org_addresses.append({
                "type": addr.get_address_type_display(),
                "text": ", ".join(parts),
            })
            if not org_iec_code and addr.iec_code:
                org_iec_code = addr.iec_code
            if not org_tax_info and addr.tax_type and addr.tax_code:
                org_tax_info = f"{addr.tax_type}: {addr.tax_code}"

    cin = ""
    if org:
        for addr in org.addresses.all():
            if addr.tax_type and "CIN" in addr.tax_type.upper():
                cin = addr.tax_code
                break

    is_draft = coa.status != APPROVED

    # -------------------------------------------------------------------------
    # Footer repeats every page (matches the PDF's per-page canvas _on_page
    # callback: hairline rule + disclaimer + "Page N", no total page count).
    # There is NO repeating per-page title header in the real PDF — the
    # company header / doc title are ordinary flowables printed once at the
    # top of the story, so we do the same below instead of using a header.
    # -------------------------------------------------------------------------
    add_page_footer(section, with_total=False)

    if is_draft:
        add_watermark(document, "DRAFT")

    # -------------------------------------------------------------------------
    # 1. Company header block — printed once, mirrors _build_header_flowables():
    #    company name (18pt bold centered) -> CIN (8pt centered) ->
    #    one line per address (8pt centered, "<b>Type Address:</b> text") ->
    #    IEC/GSTIN line (8pt centered) -> bold horizontal rule.
    # -------------------------------------------------------------------------
    p_company = document.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_company, org_name, font=FONT_HEADING, size=18, bold=True)

    if cin:
        p_cin = document.add_paragraph()
        p_cin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_cin, f"CIN: {cin}", font=FONT_BODY, size=SIZE_SMALL)

    for addr_info in org_addresses:
        p_addr = document.add_paragraph()
        p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_addr, f"{addr_info['type']} Address: ", font=FONT_BODY, size=SIZE_SMALL, bold=True)
        add_run(p_addr, addr_info["text"], font=FONT_BODY, size=SIZE_SMALL)

    iec_gstin_parts = []
    if org_iec_code:
        iec_gstin_parts.append(f"IEC Code: {org_iec_code}")
    if org_tax_info:
        iec_gstin_parts.append(org_tax_info)
    if iec_gstin_parts:
        p_iec = document.add_paragraph()
        p_iec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_iec, "    ".join(iec_gstin_parts), font=FONT_BODY, size=SIZE_SMALL)

    # Bold horizontal rule (mirrors the PDF's LINEABOVE separator table)
    rule = document.add_paragraph()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # -------------------------------------------------------------------------
    # 2. Document title — mirrors story.append(Paragraph("Certificate of Analysis", style_title))
    # -------------------------------------------------------------------------
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_title, "Certificate of Analysis", font=FONT_HEADING, size=13, bold=True)

    # -------------------------------------------------------------------------
    # 3/4. COA info rows — batch, dates, quantities (mirrors info_table)
    # -------------------------------------------------------------------------
    pg = coa.product_grade
    product_name = pg.product.name if pg else ""
    grade_str = pg.grade if pg else ""
    customer_name = coa.customer.name if coa.customer else ""

    try:
        vol_str = f"{coa.package_volume.normalize():f}".rstrip("0").rstrip(".")
    except Exception:
        vol_str = str(coa.package_volume or "")
    uom_abbr = coa.package_uom.abbreviation if coa.package_uom else ""
    pkg_type_name = coa.package_type.name if coa.package_type else ""
    supplied_qty = f"{coa.package_count} x {vol_str} {uom_abbr} {pkg_type_name}"

    date_despatch = _fmt_date(coa.date_of_despatch) if coa.date_of_despatch else "XXXX"

    pl_ci_rows = []
    if coa.packing_list_id:
        pl_ci_rows.append(("Packing List No.", _safe(coa.packing_list.pl_number)))
        try:
            ci_num = coa.packing_list.commercial_invoice.ci_number
            if ci_num:
                pl_ci_rows.append(("Commercial Invoice No.", _safe(ci_num)))
        except Exception:
            pass

    info_rows = [
        ("COA No.", _safe(coa.coa_number)),
        ("Product / Grade", f"{product_name} / {grade_str}"),
        ("Customer", customer_name),
        ("Batch No.", _safe(coa.batch_number)),
        ("Supplied Quantity", supplied_qty),
        ("Date of Despatch", date_despatch),
        ("Date of Manufacturing", _fmt_date(coa.date_of_manufacture)),
        ("Date of Retest", _fmt_date(coa.date_of_retest)),
        ("Date and Time of Sampling", _fmt_datetime(coa.date_time_of_sampling)),
        ("Date and Time of Analysis", _fmt_datetime(coa.date_time_of_analysis)),
        *pl_ci_rows,
    ]

    # Direct transliteration of:
    #   info_data = [[Paragraph(f"<b>{label}</b>", style_label), Paragraph(value, style_text)] ...]
    #   info_table = Table(info_data, colWidths=[PAGE_W*0.42, PAGE_W*0.58])
    #   info_table.setStyle(TableStyle(_GRID_STYLE))   # plain white grid, no navy bg
    info_grid_rows = [
        [{"html": f"<b>{label}</b>", "size": SIZE_BODY}, {"html": value, "size": SIZE_BODY}]
        for label, value in info_rows
    ]
    label_w = Emu(int(CONTENT_W * 0.42))
    value_w = Emu(int(CONTENT_W * 0.58))
    build_grid_table(document, info_grid_rows, [label_w, value_w])
    document.add_paragraph()

    # -------------------------------------------------------------------------
    # 5. Test parameters table — built manually (not via build_grid_table) to
    # reproduce the exact ReportLab TableStyle: a full-width navy section-title
    # row (SPAN across all cols), a navy column-header row, then data rows with
    # alternating light-grey shading on even indices (2, 4, 6 ...) exactly like
    # tbl_style's `for i in range(2, len(table_data)): if i % 2 == 0: ...`.
    # -------------------------------------------------------------------------
    col_widths = [
        Emu(int(CONTENT_W * 0.05)),   # S.No
        Emu(int(CONTENT_W * 0.22)),   # Characteristic
        Emu(int(CONTENT_W * 0.08)),   # Unit
        Emu(int(CONTENT_W * 0.10)),   # Spec Min
        Emu(int(CONTENT_W * 0.10)),   # Spec Max
        Emu(int(CONTENT_W * 0.15)),   # Spec Description
        Emu(int(CONTENT_W * 0.15)),   # Results
        Emu(int(CONTENT_W * 0.15)),   # Test Method
    ]

    col_headers = ["Sr", "Characteristic", "Unit", "Spec Min", "Spec Max",
                   "Specification", "Results", "Test Method"]

    params = list(coa.parameters.select_related("unit", "parameter", "test_method").all())
    data_rows = []
    for p in params:
        unit_str = p.unit.abbreviation if p.unit else "–"
        if p.spec_type == "QUANTITATIVE":
            spec_min_str = _esc(p.spec_min) if p.spec_min else "–"
            spec_max_str = _esc(p.spec_max) if p.spec_max else "–"
            spec_desc_str = ""
            result_str = _esc(p.result_value) if p.result_value else ""
        else:
            spec_min_str = ""
            spec_max_str = ""
            spec_desc_str = _esc(p.spec_description)
            result_str = _esc(p.result_text)

        method_str = p.test_method.code if p.test_method else ""

        data_rows.append([
            str(p.s_no),
            _safe(p.parameter.name if p.parameter else ""),
            unit_str,
            spec_min_str,
            spec_max_str,
            spec_desc_str,
            result_str,
            method_str,
        ])

    n_cols = len(col_headers)
    total_rows = 2 + len(data_rows)  # section title row + column header row + data rows
    param_table = document.add_table(rows=total_rows, cols=n_cols)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: full-width section title "TEST RESULTS" — SPAN (0,0)-(-1,0), navy bg, white bold centered
    title_cell = param_table.cell(0, 0)
    param_table.cell(0, 0).merge(param_table.cell(0, n_cols - 1))
    _set_cell_background(title_cell, COLOR_NAVY_HEX)
    set_cell_html(title_cell, "<b>TEST RESULTS</b>", font=FONT_HEADING, size=SIZE_TABLE_HEADER,
                  color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1: column header row — navy bg, white bold centered
    for c, h in enumerate(col_headers):
        cell = param_table.cell(1, c)
        _set_cell_background(cell, COLOR_NAVY_HEX)
        set_cell_html(cell, f"<b>{h}</b>", font=FONT_HEADING, size=SIZE_TABLE_HEADER,
                      color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows — alternating light-grey shading starting at first data row,
    # i.e. table_data index 2 in the PDF's flattened list (section row=0, header row=1).
    for row_idx, row in enumerate(data_rows):
        table_row_i = 2 + row_idx
        table_data_i = 2 + row_idx  # same index space as the PDF's table_data list
        shade = (table_data_i % 2 == 0)
        for c, value in enumerate(row):
            cell = param_table.cell(table_row_i, c)
            if shade:
                _set_cell_background(cell, LIGHT_GREY_HEX)
            set_cell_html(cell, value, font=FONT_BODY, size=SIZE_BODY, color=COLOR_TEXT_DARK)

    set_table_borders(param_table)
    set_column_widths(param_table, col_widths)
    document.add_paragraph()

    # -------------------------------------------------------------------------
    # 6. Closing note — mirrors story.append(Paragraph("Based on factory results", style_sig))
    # -------------------------------------------------------------------------
    p_note = document.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_note, "Based on factory results", font=FONT_BODY, size=SIZE_BODY)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
