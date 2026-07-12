"""
Commercial Invoice Word (.docx) section generator — transliterates
pdf/commercial_invoice_generator.py::build_ci_story() cell-by-cell,
span-by-span, background-by-background using pdf/docx_base.py's
build_grid_table() (the direct analog of ReportLab's Table + TableStyle).

Formatting helpers (money/date/qty formatting, CIF override math, HTML cell
strings) are reused verbatim from pdf/commercial_invoice_generator.py so the
Word output always matches the PDF's numbers — never duplicate that logic
here.

Constraint #9: this module never writes to disk; it only appends content to
an in-memory python-docx Document that the caller saves to a BytesIO buffer.
"""
import html
from decimal import Decimal

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Mm

from pdf.commercial_invoice_generator import (
    _ALL_CI_FIELDS,
    _INCOTERM_VISIBLE_FIELDS,
    _amount_to_words,
    _fmt_money,
    _fmt_qty,
    _fmt_rate,
    _party_cell_html,
    fmt_date,
    safe,
)
from pdf.utils import weight_unit_for_packing_list
from pdf.docx_base import (
    CONTENT_W,
    add_title_block,
    build_grid_table as _build_grid_table,
    build_items_table,
    build_lined_box as _build_lined_box,
)


def _unescape_html_entities(spec):
    """
    The ported ReportLab strings use HTML entities (&amp;, &nbsp;, etc.) because
    ReportLab's Paragraph is a mini HTML/XML renderer. docx_base's add_html_runs
    only understands <b>/</b>/<br/> tags — it does not decode entities — so they
    must be unescaped here before being handed to build_grid_table, otherwise
    "&amp;" renders literally instead of as "&".
    """
    if spec is None:
        return None
    if isinstance(spec, str):
        return html.unescape(spec)
    if isinstance(spec, dict) and "html" in spec:
        spec = dict(spec)
        spec["html"] = html.unescape(spec["html"])
    return spec


def build_grid_table(document, rows, col_widths, spans=None):
    """Thin wrapper around docx_base.build_grid_table that unescapes HTML
    entities in every cell before rendering (see _unescape_html_entities)."""
    clean_rows = [[_unescape_html_entities(cell) for cell in row] for row in rows]
    return _build_grid_table(document, clean_rows, col_widths, spans=spans)


def build_lined_box(document, lines_html, bg_hex=None, size=None):
    """Thin wrapper around docx_base.build_lined_box that unescapes HTML
    entities in every line before rendering (see _unescape_html_entities)."""
    clean_lines = [html.unescape(line) for line in lines_html]
    kwargs = {}
    if bg_hex is not None:
        kwargs["bg_hex"] = bg_hex
    if size is not None:
        kwargs["size"] = size
    return _build_lined_box(document, clean_lines, **kwargs)


def _exp_cell_html(label: str, exp, addr) -> str:
    """Mirrors build_ci_story()'s local _exp_cell() closure exactly."""
    parts = []
    if exp:
        parts.append(safe(getattr(exp, "name", "")))
    if addr:
        if getattr(addr, "line1", ""):
            parts.append(addr.line1)
        if getattr(addr, "line2", ""):
            parts.append(addr.line2)
        city_parts = [v for v in (getattr(addr, "city", ""), getattr(addr, "pin", ""),
                                  getattr(addr, "state", "")) if v]
        if getattr(addr, "country", None):
            city_parts.append(addr.country.name)
        if city_parts:
            parts.append(", ".join(city_parts))
        phone_cc = getattr(addr, "phone_country_code", "")
        phone_no = getattr(addr, "phone_number", "")
        if phone_cc and phone_no:
            parts.append(f"Ph: {phone_cc} {phone_no}")
        elif phone_no:
            parts.append(f"Ph: {phone_no}")
        if getattr(addr, "email", ""):
            parts.append(addr.email)
        iec = getattr(addr, "iec_code", "")
        if iec:
            parts.append(f"IEC: {iec}")
        tax_type = getattr(addr, "tax_type", "")
        tax_code_v = getattr(addr, "tax_code", "")
        if tax_type and tax_code_v:
            parts.append(f"{tax_type}: {tax_code_v}")
    body = "<br/>".join(parts)
    return f"<b>{label}</b><br/>{body}" if body else f"<b>{label}</b>"


def build_ci_word_section(document, ci, client_invoice=False, pi=None):
    """
    Append the Commercial Invoice section content to `document`, transliterating
    build_ci_story()'s Table()/TableStyle() calls line-by-line.

    Args:
        document: python-docx Document being built.
        ci: CommercialInvoice model instance.
        client_invoice: when True, line item rates are CIF-adjusted using the
            linked PI's freight/insurance_amount (same formula as the PDF/UI).
        pi: linked ProformaInvoice (currency source); resolved from the PL if
            not passed explicitly.
    """
    pl = getattr(ci, "packing_list", None)
    exp = getattr(pl, "exporter", None) if pl else None
    cons = getattr(pl, "consignee", None) if pl else None
    buyer = getattr(pl, "buyer", None) if pl else None
    notify_party_org = getattr(pl, "notify_party", None) if pl else None

    if pi is None:
        pi = getattr(pl, "proforma_invoice", None) if pl else None
    currency_code = pi.currency.code if pi else "USD"

    # ---- Title block (ONE-TIME, not a repeating header) --------------------
    # Mirrors: story.append(Paragraph(exp.name, style_company_header));
    #          story.append(Paragraph("COMMERCIAL INVOICE", style_title))
    add_title_block(document, safe(getattr(exp, "name", "")), ["Commercial Invoice"])

    col_4 = Emu(int(CONTENT_W / 4))
    col_3 = Emu(int(CONTENT_W / 3))
    col_2 = Emu(int(CONTENT_W / 2))

    corp_addr = (
        (exp.addresses.filter(address_type="OFFICE").first()
         or exp.addresses.filter(address_type="REGISTERED").first())
        if exp else None
    )
    reg_addr = (exp.addresses.filter(address_type="REGISTERED").first() if exp else None)
    factory_addr = (exp.addresses.filter(address_type="FACTORY").first() if exp else None)

    ref_lines = []
    if pl:
        po_no = safe(getattr(pl, "po_number", ""))
        po_date = fmt_date(getattr(pl, "po_date", None))
        lc_no = safe(getattr(pl, "lc_number", ""))
        lc_date = fmt_date(getattr(pl, "lc_date", None))
        bl_no = safe(getattr(pl, "bl_number", ""))
        bl_date = fmt_date(getattr(pl, "bl_date", None))
        so_no = safe(getattr(pl, "so_number", ""))
        so_date = fmt_date(getattr(pl, "so_date", None))
        other_ref = safe(getattr(pl, "other_references", ""))
        other_ref_date = fmt_date(getattr(pl, "other_references_date", None))
        if po_no:
            ref_lines.append(f"<b>PO No/Date:</b> {po_no}{(' / ' + po_date) if po_date else ''}")
        if lc_no:
            ref_lines.append(f"<b>LC No/Date:</b> {lc_no}{(' / ' + lc_date) if lc_date else ''}")
        if bl_no:
            ref_lines.append(f"<b>BL No/Date:</b> {bl_no}{(' / ' + bl_date) if bl_date else ''}")
        if so_no:
            ref_lines.append(f"<b>SO No/Date:</b> {so_no}{(' / ' + so_date) if so_date else ''}")
        if other_ref:
            ref_lines.append(
                f"<b>Other Ref/Date:</b> {other_ref}"
                f"{(' / ' + other_ref_date) if other_ref_date else ''}"
            )
    refs_cell_html = (
        "<b>Other References</b><br/>" + "<br/>".join(ref_lines)
        if ref_lines else "<b>Other References</b>"
    )

    pi_obj = getattr(pl, "proforma_invoice", None) if pl else None
    pi_number = safe(getattr(pi_obj, "pi_number", "")) if pi_obj else ""
    ci_number = safe(getattr(ci, "ci_number", ""))

    pi_number_with_date = pi_number
    if pi_obj:
        pi_date = getattr(pi_obj, "pi_date", None)
        if pi_date:
            pi_number_with_date = f"{pi_number} {fmt_date(pi_date)}"

    from apps.workflow.constants import APPROVED
    ci_status = getattr(ci, "status", None)
    if ci_status == APPROVED:
        ci_date_display = ""
        try:
            from apps.workflow.models import AuditLog
            approval_log = AuditLog.objects.filter(
                document_type="commercial_invoice",
                document_id=ci.id,
                action="APPROVE",
            ).order_by("-created_at").first()
            if approval_log:
                ci_date_display = fmt_date(approval_log.created_at.date())
        except Exception:
            pass
    else:
        from datetime import date
        ci_date_display = fmt_date(date.today())
    ci_number_with_date = f"{ci_number} {ci_date_display}" if ci_date_display else ci_number

    # ---- header_tbl: Exporter (navy, spans 2 cols) | PI No. | CI No. --------
    build_grid_table(
        document,
        [[
            {"html": "<b>Exporter</b>", "bg": "navy", "align": WD_ALIGN_PARAGRAPH.CENTER},
            None,  # covered by span below
            {"html": f"<b>Proforma Invoice No.</b><br/>{pi_number_with_date}",
             "bg": "navy", "align": WD_ALIGN_PARAGRAPH.CENTER},
            {"html": f"<b>Commercial Invoice No.</b><br/>{ci_number_with_date}",
             "bg": "navy", "align": WD_ALIGN_PARAGRAPH.CENTER},
        ]],
        col_widths=[col_4, col_4, col_4, col_4],
        spans=[(0, 0, 0, 1)],
    )

    # ---- exp_tbl: Corporate Office | Registered Office | (Factory) ----------
    office_cell = _exp_cell_html("Corporate Office", exp, corp_addr)
    reg_cell = _exp_cell_html("Registered Office Address", exp, reg_addr)
    factory_cell = _exp_cell_html("Factory Address", exp, factory_addr)
    if factory_addr:
        build_grid_table(
            document,
            [[{"html": office_cell}, {"html": reg_cell}, {"html": factory_cell}]],
            col_widths=[col_3, col_3, col_3],
        )
    else:
        build_grid_table(
            document,
            [[{"html": office_cell}, {"html": reg_cell}]],
            col_widths=[col_2, col_2],
        )

    # ---- party_tbl: Buyer | Consignee | (Notify Party) ----------------------
    buyer_cell = _party_cell_html("Buyer", buyer if buyer else cons)
    cons_cell = _party_cell_html("Consignee", cons)
    if notify_party_org:
        notify_cell = _party_cell_html("Notify Party", notify_party_org)
        build_grid_table(
            document,
            [[{"html": buyer_cell}, {"html": cons_cell}, {"html": notify_cell}]],
            col_widths=[col_3, col_3, col_3],
        )
    else:
        build_grid_table(
            document,
            [[{"html": buyer_cell}, {"html": cons_cell}]],
            col_widths=[col_2, col_2],
        )

    # ---- shipping_tbl: 4 cols x 2 rows --------------------------------------
    pre_carriage_obj = getattr(pl, "pre_carriage_by", None) if pl else None
    pre_carriage_val = safe(getattr(pre_carriage_obj, "name", "")) if pre_carriage_obj else ""
    place_receipt_obj = getattr(pl, "place_of_receipt_by_pre_carrier", None) if pl else None
    place_receipt_val = safe(getattr(place_receipt_obj, "name", "")) if place_receipt_obj else ""
    port_loading_obj = getattr(pl, "port_of_loading", None) if pl else None
    port_loading_val = safe(getattr(port_loading_obj, "name", "")) if port_loading_obj else ""
    port_discharge_obj = getattr(pl, "port_of_discharge", None) if pl else None
    port_discharge_val = safe(getattr(port_discharge_obj, "name", "")) if port_discharge_obj else ""
    final_dest_obj = getattr(pl, "final_destination", None) if pl else None
    final_dest_val = safe(getattr(final_dest_obj, "name", "")) if final_dest_obj else ""
    dest_country_obj = getattr(pl, "country_of_final_destination", None) if pl else None
    dest_country_val = safe(getattr(dest_country_obj, "name", "")) if dest_country_obj else ""
    origin_country_obj = getattr(pl, "country_of_origin", None) if pl else None
    origin_country_val = safe(getattr(origin_country_obj, "name", "")) if origin_country_obj else ""
    vessel_flight_val = safe(getattr(pl, "vessel_flight_no", "")) if pl else ""

    build_grid_table(
        document,
        [
            [
                {"html": f"<b>Pre-carriage by</b><br/>{pre_carriage_val}"},
                {"html": f"<b>Place of Receipt by Pre-Carrier</b><br/>{place_receipt_val}"},
                {"html": f"<b>Port of Loading</b><br/>{port_loading_val}"},
                {"html": f"<b>Port of Discharge</b><br/>{port_discharge_val}"},
            ],
            [
                {"html": f"<b>Final Destination</b><br/>{final_dest_val}"},
                {"html": f"<b>Country of Final Destination</b><br/>{dest_country_val}"},
                {"html": f"<b>Country of Origin of Goods</b><br/>{origin_country_val}"},
                {"html": f"<b>Vessel / Flight No.</b><br/>{vessel_flight_val}"},
            ],
        ],
        col_widths=[col_4, col_4, col_4, col_4],
    )

    # ---- terms_tbl: Payment Terms | Incoterms | Other References -----------
    incoterm_obj = getattr(pl, "incoterms", None) if pl else None
    incoterm_str = safe(getattr(incoterm_obj, "code", "")) if incoterm_obj else ""
    payment_term_obj = getattr(pl, "payment_terms", None) if pl else None
    payment_term_str = safe(getattr(payment_term_obj, "name", "")) if payment_term_obj else ""

    build_grid_table(
        document,
        [[
            {"html": f"<b>Payment Terms</b><br/>{payment_term_str}"},
            {"html": f"<b>Incoterms</b><br/>{incoterm_str}"},
            {"html": refs_cell_html},
        ]],
        col_widths=[col_3, col_3, col_3],
    )
    document.add_paragraph()

    # ---- Line items table ---------------------------------------------------
    cif_overrides = {}
    if client_invoice and pi is not None:
        all_ci_items = list(ci.line_items.all().order_by("id"))
        total_qty_all = sum(float(getattr(it, "total_quantity", 0) or 0) for it in all_ci_items)
        total_fob_all = sum(float(getattr(it, "amount", 0) or 0) for it in all_ci_items)
        total_freight = float(getattr(pi, "freight", None) or 0)
        total_insurance = float(getattr(pi, "insurance_amount", None) or 0)
        freight_per_unit = total_freight / total_qty_all if total_qty_all > 0 else 0
        for it in all_ci_items:
            base_rate = float(getattr(it, "rate", None) or 0)
            qty = float(getattr(it, "total_quantity", None) or 0)
            ins_per_unit = (total_insurance * base_rate) / total_fob_all if total_fob_all > 0 else 0
            cif_rate = base_rate + freight_per_unit + ins_per_unit
            cif_overrides[it.pk] = (cif_rate, Decimal(str(round(cif_rate * qty, 2))))

    headers = ["Sr.", "HSN Code", "No & Kind of Packages", "Item Code", "Description of Goods",
               "Qty", f"Rate ({currency_code})", f"Amount ({currency_code})"]
    rows = []
    line_items_total = Decimal("0.00")

    idx = 0
    for it in ci.line_items.all().order_by("id"):
        idx += 1
        pkg_text = safe(getattr(it, "packages_kind", ""))
        uom_obj = getattr(it, "uom", None)
        uom_display = safe(getattr(uom_obj, "abbreviation", "")) if uom_obj else ""
        qty_val = getattr(it, "total_quantity", None)
        if client_invoice and it.pk in cif_overrides:
            rate_val, amount_val = cif_overrides[it.pk]
        else:
            rate_val = getattr(it, "rate", None)
            amount_val = getattr(it, "amount", None)

        if amount_val is not None:
            try:
                line_items_total += Decimal(str(amount_val))
            except Exception:
                pass

        rows.append([
            str(idx),
            safe(it.hsn_code),
            pkg_text,
            safe(it.item_code),
            safe(it.description),
            f"{_fmt_qty(qty_val)} {uom_display}".strip(),
            _fmt_rate(rate_val),
            _fmt_money(amount_val),
        ])

    # Sr(10) + HSN(22) + Packages(24) + ItemCode(20) + Desc(46) + Qty(18) + Rate(20) + Amt(20) = 180mm
    col_widths = [Mm(10), Mm(22), Mm(24), Mm(20), Mm(46), Mm(18), Mm(20), Mm(20)]
    build_items_table(document, headers, rows, col_widths, right_cols=[6, 7])
    document.add_paragraph()

    # ---- Weights + cost breakdown -------------------------------------------
    total_net_val = Decimal("0.000")
    total_gross_val = Decimal("0.000")
    if pl:
        try:
            for cont in pl.containers.all().order_by("id"):
                for item in cont.items.all():
                    net_mat_wt = getattr(item, "net_material_weight", None)
                    if net_mat_wt is not None:
                        try:
                            total_net_val += Decimal(str(net_mat_wt))
                        except Exception:
                            pass
                gw = getattr(cont, "gross_weight", None)
                if gw is not None:
                    try:
                        total_gross_val += Decimal(str(gw))
                    except Exception:
                        pass
        except Exception:
            pass

    lc_details_val = safe(getattr(ci, "lc_details", ""))
    visible_fields = _INCOTERM_VISIBLE_FIELDS.get(incoterm_str, _ALL_CI_FIELDS)

    freight_amount = Decimal("0.00")
    insurance_amount = Decimal("0.00")
    show_freight = False
    show_insurance = False
    if not client_invoice:
        if "freight" in visible_fields and getattr(ci, "freight", None) is not None:
            try:
                freight_amount = Decimal(str(ci.freight))
                show_freight = True
            except Exception:
                pass
        if "insurance" in visible_fields and getattr(ci, "insurance", None) is not None:
            try:
                insurance_amount = Decimal(str(ci.insurance))
                show_insurance = True
            except Exception:
                pass
    invoice_total = line_items_total if client_invoice else line_items_total + freight_amount + insurance_amount

    weight_unit = weight_unit_for_packing_list(pl) if pl else "KGS"

    # ---- totals_charges_tbl: weights (left) | cost breakdown (right) --------
    left_lines = [
        f"<b>Total Net Weight:</b> {_fmt_qty(total_net_val)} {weight_unit}",
        f"<b>Total Gross Weight:</b> {_fmt_qty(total_gross_val)} {weight_unit}",
    ]
    if lc_details_val:
        left_lines.append(f"<b>L/C Details:</b> {lc_details_val}")
    left_html = "<br/>".join(left_lines)

    breakdown_header = f"COST BREAKDOWN ({incoterm_str})" if incoterm_str else "COST BREAKDOWN"
    breakdown_lines = [f"<b>{breakdown_header}</b>"]
    if incoterm_str != "EXW":
        breakdown_lines.append(f"FOB Value (Line Items): {currency_code} {_fmt_money(line_items_total)}")
    if show_freight:
        breakdown_lines.append(f"Freight: {currency_code} {_fmt_money(freight_amount)}")
    if show_insurance:
        breakdown_lines.append(f"Insurance Amount: {currency_code} {_fmt_money(insurance_amount)}")
    breakdown_html = "<br/>".join(breakdown_lines)

    build_grid_table(
        document,
        [[{"html": left_html}, {"html": breakdown_html}]],
        col_widths=[col_2, col_2],
    )

    # ---- invoice_total_tbl (navy) --------------------------------------------
    invoice_total_label = "Total CIF Amount (Payable)" if client_invoice else "Invoice Total (Amount Payable)"
    build_grid_table(
        document,
        [[
            {"html": f"<b>{invoice_total_label}</b>", "bg": "navy"},
            {"html": f"<b>{currency_code} {_fmt_money(invoice_total)}</b>", "bg": "navy",
             "align": WD_ALIGN_PARAGRAPH.RIGHT},
        ]],
        col_widths=[Mm(140), Mm(40)],
    )
    document.add_paragraph()

    # ---- Amount in words -----------------------------------------------------
    amount_in_words_str = _amount_to_words(invoice_total, currency=currency_code)
    if amount_in_words_str:
        build_grid_table(
            document,
            [[{"html": f"<b>Amount in Words:</b> {amount_in_words_str}",
               "align": WD_ALIGN_PARAGRAPH.CENTER}]],
            col_widths=[CONTENT_W],
        )
        document.add_paragraph()

    # ---- Declaration -----------------------------------------------------------
    build_grid_table(
        document,
        [[{"html": "<b>Declaration:</b> We declare that this invoice shows actual price of the "
                   "goods described and that all particulars are true and correct."}]],
        col_widths=[CONTENT_W],
    )
    document.add_paragraph()

    # ---- Banking box -----------------------------------------------------------
    # Transliterates build_ci_story()'s bank_lines list verbatim (lines ~806-834
    # of pdf/commercial_invoice_generator.py): a plain white, single-column,
    # one-line-per-row bordered box — NOT a navy title bar / 2-column grid.
    bank = getattr(ci, "bank", None)
    if bank:
        bank_lines = []
        bank_lines.append(f"<b>BENEFICIARY NAME:</b> {safe(getattr(bank, 'beneficiary_name', ''))}")
        bank_lines.append(f"<b>BANK NAME:</b> {safe(getattr(bank, 'bank_name', ''))}")
        bank_lines.append(f"<b>BRANCH NAME:</b> {safe(getattr(bank, 'branch_name', ''))}")
        bank_lines.append(f"<b>BRANCH ADDRESS:</b> {safe(getattr(bank, 'branch_address', ''))}")
        bank_lines.append(f"<b>A/C NO.:</b> {safe(getattr(bank, 'account_number', ''))}")
        if getattr(bank, "routing_number", None):
            bank_lines.append(f"<b>IFSC CODE:</b> {safe(bank.routing_number)}")
        if getattr(bank, "swift_code", None):
            bank_lines.append(f"<b>SWIFT CODE:</b> {safe(bank.swift_code)}")
        if getattr(bank, "iban", None):
            bank_lines.append(f"<b>IBAN:</b> {safe(bank.iban)}")
        if safe(getattr(bank, "intermediary_bank_name", "")):
            intermediary_currency_code = (
                safe(getattr(bank.intermediary_currency, "code", ""))
                if getattr(bank, "intermediary_currency", None) else ""
            )
            bank_lines.append(
                f"<b>Intermediary Institution Routing for Currency</b> {intermediary_currency_code} "
                f"<b>A/C No.:</b> {safe(bank.intermediary_account_number)} "
                f"The Bank of {safe(bank.intermediary_bank_name)} "
                f"<b>SWIFT Code:</b> {safe(bank.intermediary_swift_code)}"
            )
        bank_lines.append(
            "Request your bank to send MT 103 Message to our bank and send us copy of this "
            "message to trace &amp; claim the payment from our bank."
        )
        build_lined_box(document, bank_lines)
