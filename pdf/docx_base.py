"""
Shared design-system constants and drawing utilities for TradeDocs Word (.docx)
generators.

IMPORTANT: these mirror the styling ACTUALLY used by the ReportLab PDF
generators (pdf/proforma_invoice.py, pdf/purchase_order.py, pdf/packing_list*.py,
pdf/commercial_invoice_generator.py, pdf/certificate_of_analysis.py) — navy
header bars (#1A2B4B) with white bold centered text, Helvetica body text, a
red diagonal DRAFT watermark, and a per-page footer only (no repeating
per-page title header — the company name / document title are ordinary
flowables printed once at the top of the content, exactly like the PDF's
`story.append(Paragraph(...))` calls). pdf/base.py's black-and-white palette
is NOT used by any real generator — do not copy it.
"""
import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

# ---- Colour palette (matches the real PDF generators) ------------------------
COLOR_NAVY = RGBColor(0x1A, 0x2B, 0x4B)
COLOR_NAVY_HEX = "1A2B4B"
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT_DARK = RGBColor(0x00, 0x00, 0x00)
COLOR_BORDER = "000000"
# Simulated "red at 15% opacity over white" (reportlab draws #CC0000 alpha .15)
COLOR_WATERMARK = RGBColor(0xF7, 0xD9, 0xD9)

# ---- Fonts (Arial is the Word-native metric-compatible equivalent of the
# PDF's Helvetica; the PDF never uses Times) -----------------------------------
FONT_HEADING = "Arial"      # Helvetica-Bold: company name, document title
FONT_BODY = "Arial"         # Helvetica: body text, data cells
FONT_MONO = "Courier New"

# ---- Type sizes (pt) — match the values hardcoded in the PDF generators ------
SIZE_COMPANY = 18
SIZE_DOC_TITLE = 13
SIZE_TABLE_HEADER = 9
SIZE_BODY = 9
SIZE_SMALL = 8

# ---- Page geometry (A4, matches every PDF generator's margins) --------------
PAGE_W = Mm(210)
PAGE_H = Mm(297)
MARGIN_H = Mm(15)
MARGIN_TOP = Mm(15)
MARGIN_BOTTOM = Mm(20)
CONTENT_W = PAGE_W - 2 * MARGIN_H


def setup_page(document):
    """Apply A4 size + margins matching the PDF layout to every section."""
    # python-docx's default template bakes 10pt space-after + 1.15x line spacing
    # into every paragraph (word/styles.xml docDefaults) — reportlab's PDFs have
    # no such padding, so leaving this in place makes every table cell taller in
    # Word than in the PDF and pushes content onto extra pages. Zero it here,
    # once, since every generator calls setup_page() right after Document().
    normal = document.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for section in document.sections:
        section.page_width = PAGE_W
        section.page_height = PAGE_H
        section.left_margin = MARGIN_H
        section.right_margin = MARGIN_H
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.header_distance = Mm(8)
        section.footer_distance = Mm(10)


# ---- Low-level oxml helpers ---------------------------------------------------

def _set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color=COLOR_BORDER, sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_table_borders(table, color=COLOR_BORDER, sz=6):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _set_cell_valign(cell, valign="top"):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), valign)
    tcPr.append(va)


def set_column_widths(table, widths):
    """widths: list of Length values (e.g. Mm(...)/Emu(...)), one per column."""
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
    for column, width in zip(table.columns, widths):
        column.width = width


# ---- Paragraph / run helpers ---------------------------------------------------

def add_run(paragraph, text, font=FONT_BODY, size=SIZE_BODY, bold=False,
            italic=False, color=COLOR_TEXT_DARK, align=None):
    run = paragraph.add_run(str(text if text not in (None, "") else ""))
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if align is not None:
        paragraph.alignment = align
    return run


# Tiny subset of ReportLab's Paragraph markup (<b>...</b>, <br/>) — the real
# generators build almost every cell as Paragraph(f"<b>Label</b><br/>{value}",
# style), so parsing this markup lets word generators port that content
# nearly verbatim instead of re-deriving structure.
_TAG_RE = re.compile(r"(<b>|</b>|<br\s*/?>)", re.IGNORECASE)


def add_html_runs(paragraph, html, font=FONT_BODY, size=SIZE_BODY, color=COLOR_TEXT_DARK):
    bold = False
    for part in _TAG_RE.split(html or ""):
        if not part:
            continue
        low = part.lower()
        if low == "<b>":
            bold = True
        elif low == "</b>":
            bold = False
        elif low.startswith("<br"):
            paragraph.add_run().add_break()
        else:
            run = paragraph.add_run(part)
            run.font.name = font
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color


def set_cell_html(cell, html, font=FONT_BODY, size=SIZE_BODY, color=COLOR_TEXT_DARK,
                   align=None, valign="top"):
    """Replace a cell's content with parsed <b>/<br/> markup, ReportLab-Paragraph-style."""
    p = cell.paragraphs[0]
    p.text = ""
    if align is not None:
        p.alignment = align
    add_html_runs(p, html, font=font, size=size, color=color)
    _set_cell_valign(cell, valign)
    return p


def cell_text(cell, text, font=FONT_BODY, size=SIZE_BODY, bold=False,
              color=COLOR_TEXT_DARK, align=None, valign="top"):
    """Replace a table cell's content with a single styled run (no markup)."""
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    if align is not None:
        paragraph.alignment = align
    add_run(paragraph, text, font=font, size=size, bold=bold, color=color)
    _set_cell_valign(cell, valign)
    return paragraph


def _add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ---- One-time title block (NOT a repeating header) ----------------------------
#
# None of the real PDF generators draw a per-page canvas header — the company
# name and document title are ordinary flowables at the top of the content
# (story.append(Paragraph(name, style_company_header)) etc.), so they only
# appear once, at the very top, exactly like calling this from body content.

def add_title_block(document, company_name, doc_title_lines):
    """
    company_name    : str — printed once, centered, bold, 18pt
    doc_title_lines : list[str] — printed once below the name, centered, bold, 13pt,
                      followed by a bold horizontal rule (mirrors the PDF's
                      LINEABOVE separator table).
    """
    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, company_name or "", font=FONT_HEADING, size=SIZE_COMPANY, bold=True)

    for line in (doc_title_lines or []):
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, str(line).upper(), font=FONT_HEADING, size=SIZE_DOC_TITLE, bold=True)

    rule = document.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---- Footer (repeats every page, matches the PDF's per-page canvas footer) ---

def add_page_footer(section, with_total=True):
    """
    Disclaimer + page number, matching draw_footer()/_draw_footer() in the PDF
    generators. Pass with_total=False for generators whose PDF only prints
    "Page N" (e.g. COA, the CI/PL sub-generators) rather than "Page N of Total".
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "This is a computer generated document and does not require signature",
            size=8, color=COLOR_TEXT_DARK)

    p_page = footer.add_paragraph()
    p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_page, "Page ", size=7)
    _add_field(p_page, "PAGE")
    if with_total:
        add_run(p_page, " of ", size=7)
        _add_field(p_page, "NUMPAGES")


def add_watermark(document, text="DRAFT"):
    """
    Faint diagonal DRAFT watermark on every page, via a rotated DrawingML
    (WordprocessingShape) textbox injected into each section's header — there
    is no native python-docx watermark API. DrawingML is used instead of the
    legacy VML watermark recipe because VML shape `rotation` is unreliably
    rendered by common OOXML consumers; DrawingML's `a:xfrm rot=` is the
    modern, reliably-rendered mechanism. Rotated -45 degrees, 80pt bold,
    colored to simulate the PDF's red-at-15%-opacity look.
    """
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ""
        run = p.add_run()
        color_hex = "%02X%02X%02X" % (COLOR_WATERMARK[0], COLOR_WATERMARK[1], COLOR_WATERMARK[2])
        # rot is in 60,000ths of a degree, clockwise. 315 degrees == -45 degrees.
        xml = (
            '<w:drawing '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
            '<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0" simplePos="0" '
            'locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="2">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="page"><wp:align>center</wp:align></wp:positionH>'
            '<wp:positionV relativeFrom="page"><wp:align>center</wp:align></wp:positionV>'
            '<wp:extent cx="5486400" cy="1828800"/>'
            '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            '<wp:wrapNone/>'
            '<wp:docPr id="1" name="DraftWatermark"/>'
            '<wp:cNvGraphicFramePr/>'
            '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
            '<wps:wsp>'
            '<wps:cNvSpPr txBox="1"/>'
            '<wps:spPr>'
            '<a:xfrm rot="18900000"><a:off x="0" y="0"/><a:ext cx="5486400" cy="1828800"/></a:xfrm>'
            '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
            '<a:noFill/><a:ln><a:noFill/></a:ln>'
            '</wps:spPr>'
            '<wps:txbx><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
            f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:b/>'
            f'<w:color w:val="{color_hex}"/><w:sz w:val="160"/></w:rPr><w:t>{text}</w:t></w:r>'
            '</w:p></w:txbxContent></wps:txbx>'
            '<wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr">'
            '<a:noAutofit/></wps:bodyPr>'
            '</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>'
        )
        drawing = parse_xml(xml)
        run._r.append(drawing)


# ---- Generic grid table builder (mirrors reportlab Table()+TableStyle()) -----

def build_grid_table(document, rows, col_widths, spans=None):
    """
    Generic bordered table mirroring how the PDF generators build Table(data)
    + TableStyle([...]) calls, so a doc-specific word generator can transliterate
    a PDF table almost line-by-line.

    rows      : list of list of cell specs, one inner list per row. Each spec is
                either a plain string/html snippet (defaults applied), or a dict:
                {"html": "<b>Label</b><br/>value", "bg": "navy"|None,
                 "color": RGBColor (optional override), "align": WD_ALIGN_PARAGRAPH.*,
                 "font": str, "size": int}
                Use None for a cell that will be covered by a merge from `spans`.
    col_widths: list of Length values, one per column.
    spans     : list of (r1, c1, r2, c2) 0-indexed inclusive merge ranges.
    """
    n_rows = len(rows)
    n_cols = len(col_widths)
    table = document.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r, row in enumerate(rows):
        for c, spec in enumerate(row):
            if spec is None:
                continue
            if isinstance(spec, str):
                spec = {"html": spec}
            cell = table.cell(r, c)
            bg = spec.get("bg")
            if bg == "navy":
                _set_cell_background(cell, COLOR_NAVY_HEX)
                default_color = COLOR_WHITE
                default_font = FONT_HEADING
            elif bg:
                # Any other value is treated as a literal hex fill (e.g. "F5F5F5"),
                # matching a specific real generator's non-navy shaded box.
                _set_cell_background(cell, bg)
                default_color = COLOR_TEXT_DARK
                default_font = FONT_BODY
            else:
                default_color = COLOR_TEXT_DARK
                default_font = FONT_BODY
            set_cell_html(
                cell, spec.get("html", ""),
                font=spec.get("font", default_font),
                size=spec.get("size", SIZE_BODY),
                color=spec.get("color", default_color),
                align=spec.get("align"),
            )

    for (r1, c1, r2, c2) in (spans or []):
        table.cell(r1, c1).merge(table.cell(r2, c2))

    set_table_borders(table)
    set_column_widths(table, col_widths)
    return table


def build_items_table(document, headers, rows, col_widths, right_cols=None):
    """
    Styled line-items table: navy header row (white bold centered text),
    matching every real generator's item table.
    headers    : list[str]
    rows       : list[list[str]] — each inner list is one data row
    col_widths : list[Length]
    right_cols : list[int] — 0-based column indices to right-align
    """
    right_cols = set(right_cols or [])
    header_row = [
        {"html": h, "bg": "navy", "align": WD_ALIGN_PARAGRAPH.CENTER,
         "size": SIZE_TABLE_HEADER, "font": FONT_HEADING}
        for h in headers
    ]
    body_rows = []
    for row in rows:
        body_row = []
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT
            body_row.append({"html": str(value if value is not None else ""),
                              "align": align, "size": SIZE_SMALL})
        body_rows.append(body_row)
    return build_grid_table(document, [header_row] + body_rows, col_widths)


def build_lined_box(document, lines_html, bg_hex=None, size=SIZE_BODY):
    """
    Single-column bordered box, one row per line — matches the ACTUAL bank-details
    (and similar) boxes in pdf/proforma_invoice.py, pdf/commercial_invoice_generator.py,
    and pdf/purchase_order.py: a plain bordered table, each row is one inline
    "<b>Label:</b> value" line (NOT a 2-column label/value grid, and NOT a navy
    title bar — none of the real generators do either of those for this box).

    lines_html : list[str] — each already-formatted line, e.g. "<b>Bank Name:</b> HDFC"
    bg_hex     : None for plain white (PI/CI), or a literal hex string like
                 "F5F5F5" for a shaded box (PO uses light gray, not navy).
    """
    if not lines_html:
        return None
    rows = [[{"html": line, "bg": bg_hex, "size": size}] for line in lines_html]
    return build_grid_table(document, rows, [CONTENT_W])
