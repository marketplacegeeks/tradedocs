"""
Purchase Order Word (.docx) Generator

Transliterates generate_purchase_order_pdf_bytes()'s ReportLab Table()/TableStyle()
calls into docx tables using pdf/docx_base.py's build_grid_table/build_items_table
primitives — cell by cell, span by span, background by background — so the Word
output is visually consistent with the real PDF (navy header bars, Helvetica-style
body text, one-time title block, footer-only repeat, DRAFT watermark).
"""
import io
from decimal import Decimal

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Mm

from pdf.docx_base import (
    COLOR_NAVY_HEX,
    COLOR_WHITE,
    CONTENT_W,
    FONT_HEADING,
    SIZE_BODY,
    SIZE_SMALL,
    SIZE_TABLE_HEADER,
    _set_cell_background,
    add_page_footer,
    add_run,
    add_title_block,
    add_watermark,
    build_grid_table,
    build_items_table,
    build_lined_box,
    set_cell_html,
    setup_page,
)
from pdf.purchase_order import (
    _addr_lines,
    _amount_in_words,
    _fmt_cur,
    _fmt_qty,
    _org_first_address,
    _safe,
)


# ============================================================================
# MAIN GENERATOR
# ============================================================================

def generate_po_docx(po) -> io.BytesIO:
    """
    Generate a Purchase Order Word document.

    Args:
        po: PurchaseOrder model instance (with prefetched line_items)

    Returns:
        io.BytesIO: in-memory .docx content
    """
    from apps.workflow.constants import APPROVED

    document = docx.Document()
    setup_page(document)
    add_page_footer(document.sections[0], with_total=True)

    is_draft = getattr(po, "status", None) != APPROVED
    if is_draft:
        add_watermark(document, "DRAFT")

    # ========================================================================
    # SECTION 1: BUYER LETTERHEAD (shown only when a buyer org is set)
    # ========================================================================

    buyer_org = getattr(po, "buyer", None)
    if buyer_org:
        buyer_org_name = _safe(getattr(buyer_org, "name", "")).strip()
        if buyer_org_name:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, buyer_org_name, size=13, bold=True)

        try:
            buyer_reg_addr = buyer_org.addresses.filter(address_type="REGISTERED").first()
        except Exception:
            buyer_reg_addr = None

        if buyer_reg_addr:
            addr_parts = list(filter(None, [
                _safe(buyer_reg_addr.line1).strip(),
                _safe(buyer_reg_addr.line2).strip(),
                _safe(buyer_reg_addr.city).strip(),
                _safe(buyer_reg_addr.pin).strip(),
                _safe(getattr(buyer_reg_addr.country, "name", "")).strip() if getattr(buyer_reg_addr, "country", None) else "",
            ]))
            if addr_parts:
                p_addr = document.add_paragraph()
                p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p_addr, " | ".join(addr_parts), size=SIZE_BODY)

    # ========================================================================
    # SECTION 2: DOCUMENT TITLE (one-time, matches the PDF's title + rule)
    # ========================================================================

    add_title_block(document, "", ["Purchase Order"])

    # ========================================================================
    # SECTION 3: VENDOR + PO DETAILS HEADER (2x2 grid, plain white bg)
    # ========================================================================

    vendor = getattr(po, "vendor", None)
    vendor_name = _safe(getattr(vendor, "name", ""))
    vendor_addr = _org_first_address(vendor)
    vendor_addr_lines = _addr_lines(vendor_addr)
    vendor_tax_parts = []
    if vendor_addr:
        tax_type = _safe(getattr(vendor_addr, "tax_type", "")).strip()
        tax_code = _safe(getattr(vendor_addr, "tax_code", "")).strip()
        if tax_code:
            label = f"{tax_type}: " if tax_type else ""
            vendor_tax_parts.append(f"{label}{tax_code}")
    vendor_detail_html = "<br/>".join([vendor_name] + vendor_addr_lines + vendor_tax_parts) if vendor_name else "—"

    currency_obj = getattr(po, "currency", None)
    currency_code = _safe(getattr(currency_obj, "code", ""))

    payment_terms_obj = getattr(po, "payment_terms", None)
    payment_terms_name = _safe(getattr(payment_terms_obj, "name", "")) if payment_terms_obj else "—"

    country_obj = getattr(po, "country_of_origin", None)
    country_name = _safe(getattr(country_obj, "name", "")) if country_obj else "—"

    contact_name = _safe(getattr(po, "internal_contact", "")) or "—"
    buyer_contact_html = f"<b>Buyer Contact:</b> {contact_name}"

    delivery_addr = getattr(po, "delivery_address", None)
    delivery_org_name = ""
    if delivery_addr:
        delivery_org = getattr(delivery_addr, "organisation", None)
        if delivery_org:
            delivery_org_name = _safe(getattr(delivery_org, "name", ""))
    delivery_lines = _addr_lines(delivery_addr)
    if delivery_org_name:
        delivery_lines = [delivery_org_name] + delivery_lines
    if delivery_addr:
        d_tax_type = _safe(getattr(delivery_addr, "tax_type", "")).strip()
        d_tax_code = _safe(getattr(delivery_addr, "tax_code", "")).strip()
        if d_tax_code:
            d_label = f"{d_tax_type}: " if d_tax_type else ""
            delivery_lines.append(f"{d_label}{d_tax_code}")
    delivery_html = "<br/>".join(delivery_lines) if delivery_lines else "—"

    half_w = Emu(int(CONTENT_W / 2))
    header_rows = [
        [
            {"html": f"<b>Delivery Address:</b><br/>{delivery_html}"},
            {"html": f"<b>Vendor / Supplier:</b><br/>{vendor_detail_html}"},
        ],
        [
            {"html": (
                f"<b>PO No:</b> {_safe(po.po_number)}<br/>"
                f"<b>Date:</b> {_safe(po.po_date)}<br/>"
                f"<b>Customer No:</b> {_safe(po.customer_no) or '—'}"
            )},
            {"html": buyer_contact_html + f"<br/><b>Currency:</b> {currency_code or '—'}"},
        ],
    ]
    build_grid_table(document, header_rows, [half_w, half_w])

    # ========================================================================
    # SECTION 4: DOCUMENT DETAILS STRIP (3x4 grid, plain white bg, one span)
    # ========================================================================

    tx_type = _safe(getattr(po, "transaction_type", ""))
    tx_labels = {
        "IGST": "IGST (Inter-State)",
        "CGST_SGST": "CGST+SGST (Same State)",
        "ZERO_RATED": "Zero Rated (Export)",
    }
    tx_display = tx_labels.get(tx_type, tx_type or "—")

    pol_obj = getattr(po, "port_of_loading", None)
    pol_name = _safe(getattr(pol_obj, "name", "")) if pol_obj else "—"

    pod_obj = getattr(po, "port_of_discharge", None)
    pod_name = _safe(getattr(pod_obj, "name", "")) if pod_obj else "—"

    pofd_obj = getattr(po, "port_of_final_destination", None)
    pofd_name = _safe(getattr(pofd_obj, "name", "")) if pofd_obj else "—"

    pkg_obj = getattr(po, "type_of_package", None)
    pkg_name = _safe(getattr(pkg_obj, "name", "")) if pkg_obj else "—"

    partial_raw = _safe(getattr(po, "partial_shipment", "")).strip()
    partial_display = {"YES": "Yes", "NO": "No"}.get(partial_raw, "—")

    transport_instr = _safe(getattr(po, "transport_instruction", "")).strip() or "—"

    quarter_w = Emu(int(CONTENT_W / 4))
    details_rows = [
        [
            {"html": f"<b>Payment Terms:</b><br/>{payment_terms_name}"},
            {"html": f"<b>Country of Origin:</b><br/>{country_name}"},
            {"html": f"<b>Time of Delivery:</b><br/>{_safe(po.time_of_delivery) or '—'}"},
            {"html": f"<b>Transaction Type:</b><br/>{tx_display}"},
        ],
        [
            {"html": f"<b>Port of Loading:</b><br/>{pol_name}"},
            {"html": f"<b>Port of Discharge:</b><br/>{pod_name}"},
            {"html": f"<b>Port of Final Destination:</b><br/>{pofd_name}"},
            {"html": f"<b>Packaging Type:</b><br/>{pkg_name}"},
        ],
        [
            {"html": f"<b>Partial Shipment Allowed:</b><br/>{partial_display}"},
            {"html": f"<b>Transport Instruction:</b><br/>{transport_instr}"},
            None,
            None,
        ],
    ]
    build_grid_table(
        document, details_rows, [quarter_w, quarter_w, quarter_w, quarter_w],
        spans=[(2, 1, 2, 3)],
    )

    internal_contract_number = _safe(getattr(po, "internal_contract_number", "")).strip()
    if internal_contract_number:
        build_grid_table(
            document,
            [[{"html": f"<b>Internal Contract Number:</b> {internal_contract_number}"}]],
            [Emu(int(CONTENT_W))],
        )

    document.add_paragraph()

    # ========================================================================
    # SECTION 5: LINE ITEMS TABLE (navy header + navy totals row; columns vary
    # by transaction_type — must match generate_purchase_order_pdf_bytes exactly)
    # ========================================================================

    line_items = list(po.line_items.all().order_by("sort_order", "id"))

    has_item_code = any(_safe(i.item_code).strip() for i in line_items)
    has_hsn = any(_safe(i.hsn_code).strip() for i in line_items)
    has_mfr = any(_safe(i.manufacturer).strip() for i in line_items)

    cur_sym = currency_code

    if tx_type == "IGST":
        headers = ["#", "Description"]
        col_widths_mm = [8, 48]
        if has_hsn:
            headers.append("HSN Code"); col_widths_mm.append(18)
        if has_mfr:
            headers.append("Mfr"); col_widths_mm.append(16)
        headers += ["Qty", "Unit Price", "Taxable Amt", "IGST %", "IGST Amt", "Total"]
        used = sum(col_widths_mm)
        remaining = 180 - used
        n_numeric = 6
        col_widths_mm += [remaining / n_numeric] * n_numeric
        opt_start = 2 + int(has_hsn) + int(has_mfr)
    elif tx_type == "CGST_SGST":
        headers = ["#", "Description"]
        col_widths_mm = [7, 38]
        if has_hsn:
            headers.append("HSN Code"); col_widths_mm.append(15)
        if has_mfr:
            headers.append("Mfr"); col_widths_mm.append(13)
        headers += ["Qty", "Unit Price", "Taxable Amt", "CGST %", "CGST Amt", "SGST %", "SGST Amt", "Total"]
        used = sum(col_widths_mm)
        remaining = 180 - used
        n_numeric = 8
        col_widths_mm += [remaining / n_numeric] * n_numeric
        opt_start = 2 + int(has_hsn) + int(has_mfr)
    else:
        headers = ["#", "Description"]
        col_widths_mm = [8, 52]
        if has_item_code:
            headers.append("Item Code"); col_widths_mm.append(22)
        if has_hsn:
            headers.append("HSN Code"); col_widths_mm.append(18)
        if has_mfr:
            headers.append("Mfr"); col_widths_mm.append(18)
        headers += ["Qty", "Unit Price", "Total"]
        used = sum(col_widths_mm)
        remaining = 180 - used
        n_numeric = 3
        col_widths_mm += [remaining / n_numeric] * n_numeric
        opt_start = 2 + int(has_item_code) + int(has_hsn) + int(has_mfr)

    right_cols = list(range(opt_start, len(headers)))

    rows = []
    grand_total = Decimal("0.00")
    total_taxable = Decimal("0.00")
    total_igst = Decimal("0.00")
    total_cgst = Decimal("0.00")
    total_sgst = Decimal("0.00")

    for idx, item in enumerate(line_items, start=1):
        uom_obj = getattr(item, "uom", None)
        uom_display = _safe(getattr(uom_obj, "abbreviation", "")) if uom_obj else ""
        qty_str = f"{_fmt_qty(item.quantity)} {uom_display}".strip()
        item_total = item.total or Decimal("0.00")
        try:
            grand_total += Decimal(str(item_total))
            total_taxable += Decimal(str(item.taxable_amount or 0))
            if tx_type == "IGST":
                total_igst += Decimal(str(item.igst_amount or 0))
            elif tx_type == "CGST_SGST":
                total_cgst += Decimal(str(item.cgst_amount or 0))
                total_sgst += Decimal(str(item.sgst_amount or 0))
        except Exception:
            pass

        if tx_type == "IGST":
            row = [str(idx), _safe(item.description)]
            if has_hsn:
                row.append(_safe(item.hsn_code))
            if has_mfr:
                row.append(_safe(item.manufacturer))
            row += [
                qty_str,
                _fmt_cur(item.unit_price, cur_sym),
                _fmt_cur(item.taxable_amount, cur_sym),
                _safe(item.igst_percent) or "—",
                _fmt_cur(item.igst_amount, cur_sym) if item.igst_amount else "—",
                _fmt_cur(item_total, cur_sym),
            ]
        elif tx_type == "CGST_SGST":
            row = [str(idx), _safe(item.description)]
            if has_hsn:
                row.append(_safe(item.hsn_code))
            if has_mfr:
                row.append(_safe(item.manufacturer))
            row += [
                qty_str,
                _fmt_cur(item.unit_price, cur_sym),
                _fmt_cur(item.taxable_amount, cur_sym),
                _safe(item.cgst_percent) or "—",
                _fmt_cur(item.cgst_amount, cur_sym) if item.cgst_amount else "—",
                _safe(item.sgst_percent) or "—",
                _fmt_cur(item.sgst_amount, cur_sym) if item.sgst_amount else "—",
                _fmt_cur(item_total, cur_sym),
            ]
        else:
            row = [str(idx), _safe(item.description)]
            if has_item_code:
                row.append(_safe(item.item_code))
            if has_hsn:
                row.append(_safe(item.hsn_code))
            if has_mfr:
                row.append(_safe(item.manufacturer))
            row += [
                qty_str,
                _fmt_cur(item.unit_price, cur_sym),
                _fmt_cur(item_total, cur_sym),
            ]
        rows.append(row)

    col_widths = [Mm(w) for w in col_widths_mm]
    items_table = build_items_table(document, headers, rows, col_widths, right_cols=right_cols)

    # Totals row: navy background, white bold text — appended as an extra row
    # directly on the table (mirrors the PDF's single totals row with
    # BACKGROUND/TEXTCOLOR TableStyle commands on (0,-1)-(-1,-1)).
    n_cols = len(headers)
    new_row = items_table.add_row()
    totals_cells = new_row.cells
    for c in range(n_cols):
        _set_cell_background(totals_cells[c], COLOR_NAVY_HEX)
        set_cell_html(totals_cells[c], "", color=COLOR_WHITE)

    set_cell_html(totals_cells[1], "<b>Grand Total</b>", font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE)

    if tx_type == "IGST":
        set_cell_html(totals_cells[opt_start + 2], f"<b>{_fmt_cur(total_taxable, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_html(totals_cells[opt_start + 4], f"<b>{_fmt_cur(total_igst, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_html(totals_cells[opt_start + 5], f"<b>{_fmt_cur(grand_total, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    elif tx_type == "CGST_SGST":
        set_cell_html(totals_cells[opt_start + 2], f"<b>{_fmt_cur(total_taxable, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_html(totals_cells[opt_start + 4], f"<b>{_fmt_cur(total_cgst, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_html(totals_cells[opt_start + 6], f"<b>{_fmt_cur(total_sgst, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_html(totals_cells[opt_start + 7], f"<b>{_fmt_cur(grand_total, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        set_cell_html(totals_cells[opt_start + 2], f"<b>{_fmt_cur(grand_total, cur_sym)}</b>",
                      font=FONT_HEADING, size=SIZE_SMALL, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    document.add_paragraph()

    # ========================================================================
    # SECTION 6: AMOUNT IN WORDS (plain white box, full width)
    # ========================================================================

    currency_name = _safe(getattr(currency_obj, "name", "")).strip() or currency_code
    words_str = _amount_in_words(grand_total, currency_name)
    if words_str:
        build_grid_table(
            document,
            [[{"html": f"<b>Amount in Words:</b> {words_str}"}]],
            [Emu(int(CONTENT_W))],
        )
        document.add_paragraph()

    # ========================================================================
    # SECTION 7: LINE ITEM REMARKS (optional, plain white box)
    # ========================================================================

    line_item_remarks = _safe(getattr(po, "line_item_remarks", "")).strip()
    if line_item_remarks:
        build_grid_table(
            document,
            [[{"html": f"<b>Remark:</b> {line_item_remarks}"}]],
            [Emu(int(CONTENT_W))],
        )
        document.add_paragraph()

    # ========================================================================
    # SECTION 8: BANK DETAILS (optional) — single-column light-gray lined box,
    # matching pdf/purchase_order.py lines 714-741 exactly (no navy title bar).
    # ========================================================================

    bank = getattr(po, "bank", None)
    if bank:
        bank_lines = [
            f"<b>Beneficiary Name:</b> {_safe(bank.beneficiary_name)}",
            f"<b>Bank Name:</b> {_safe(bank.bank_name)}",
            f"<b>Branch:</b> {_safe(bank.branch_name)}",
            f"<b>Account No.:</b> {_safe(bank.account_number)}",
        ]
        if bank.routing_number:
            bank_lines.append(f"<b>IFSC / Routing:</b> {_safe(bank.routing_number)}")
        if bank.swift_code:
            bank_lines.append(f"<b>SWIFT Code:</b> {_safe(bank.swift_code)}")
        if bank.iban:
            bank_lines.append(f"<b>IBAN:</b> {_safe(bank.iban)}")

        build_lined_box(document, bank_lines, bg_hex="F5F5F5")
        document.add_paragraph()

    # ========================================================================
    # SECTION 9: REMARKS BELOW TOTAL (optional, plain white box)
    # ========================================================================

    remarks = _safe(getattr(po, "remarks", "")).strip()
    if remarks:
        build_grid_table(
            document,
            [[{"html": f"<b>Remark:</b> {remarks}"}]],
            [Emu(int(CONTENT_W))],
        )
        document.add_paragraph()

    # ========================================================================
    # SECTION 10: TERMS & CONDITIONS (optional, new page, navy header bar)
    # ========================================================================

    tc_content = _safe(getattr(po, "tc_content", "")).strip()
    if tc_content:
        document.add_page_break()
        build_grid_table(
            document,
            [[{"html": "<b>Terms &amp; Conditions</b>", "bg": "navy", "size": SIZE_TABLE_HEADER}]],
            [Emu(int(CONTENT_W))],
        )
        document.add_paragraph()
        from pdf.utils import strip_html
        try:
            tc_text = strip_html(tc_content)
        except Exception:
            tc_text = tc_content
        for line in tc_text.split("\n"):
            p_line = document.add_paragraph()
            add_run(p_line, line, size=SIZE_SMALL)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
